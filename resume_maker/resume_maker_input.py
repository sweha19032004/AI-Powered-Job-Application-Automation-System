import base64
import streamlit as st
import logging
import os
from resume import Resume
from resume_maker_pdfmaker import DocxToPDF
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
# Configure logging
logging.basicConfig(filename="resume_generator.log", level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

st.set_page_config(page_title="Customizable Resume Maker", page_icon=":invoice:")
st.title("Customizable ATS-Friendly Resume Generator")
st.write("Fill in the details below to generate a fully customizable resume.")

def validate_input(inputs):
    """Validate only required fields"""
    return all(inputs)

def generate_download_link(file_path):
    """Generate a download link for the generated PDF or DOCX file."""
    with open(file_path, "rb") as f:
        file_bytes = f.read()
    b64 = base64.b64encode(file_bytes).decode()
    file_name = os.path.basename(file_path)
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{file_name}">Download {file_name}</a>'

# Initialize session state for dynamic fields
if "education_entries" not in st.session_state:
    st.session_state.education_entries = []
if "experience_entries" not in st.session_state:
    st.session_state.experience_entries = []
if "project_entries" not in st.session_state:
    st.session_state.project_entries = []

# 🟢 Resume Input Form
with st.form(key="resume_form"):
    name = st.text_input("*Your Name", placeholder="Enter your full name")
    position = st.text_input("*Position/Role", placeholder="Enter the role you are applying for")

    st.subheader("Contact Information")
    email = st.text_input("*Email", placeholder="Enter your email")
    phone = st.text_input("*Phone", placeholder="Enter your phone number")
    linkedin = st.text_input("LinkedIn Profile (Optional)", placeholder="Enter LinkedIn URL")
    github = st.text_input("GitHub Profile (Optional)", placeholder="Enter GitHub URL")
    location = st.text_input("Location", placeholder="Enter your city, country")

    st.subheader("About Me")
    about = st.text_area("About Me (Optional)", placeholder="Brief summary about yourself ...")

    st.subheader("Skills & Expertise")
    skills = st.text_area("*Skills", placeholder="Programming languages, Tools, Soft skills ...")

    st.subheader("Job Description")
    job_role = st.text_input("*Job Role", placeholder="Enter the job role you are applying for")
    job_responsibilities = st.text_area("*Job Responsibilities", placeholder="Enter job responsibilities")

    submit_button = st.form_submit_button(label="Generate Resume")

# 🟢 Dynamic Fields Handling
st.subheader("Education")
if st.button("+ Add Education"):
    st.session_state.education_entries.append({"Institution": "", "Course": "", "CGPA": "", "Duration": ""})

for i, edu in enumerate(st.session_state.education_entries):
    with st.expander(f"Education {i + 1}"):
        edu["Institution"] = st.text_input(f"Institution {i + 1}", edu["Institution"], key=f"edu_institution_{i}")
        edu["Course"] = st.text_input(f"Course {i + 1}", edu["Course"], key=f"edu_course_{i}")
        edu["CGPA"] = st.text_input(f"Percentage/CGPA {i + 1}", edu["CGPA"], key=f"edu_cgpa_{i}")
        edu["Duration"] = st.text_input(f"Duration {i + 1}", edu["Duration"], key=f"edu_duration_{i}")

st.subheader("Work Experience")
if st.button("+ Add Work Experience"):
    st.session_state.experience_entries.append({"Company": "", "Job Title": "", "Duration": "", "Responsibilities": ""})

for i, exp in enumerate(st.session_state.experience_entries):
    with st.expander(f"Experience {i + 1}"):
        exp["Company"] = st.text_input(f"Company {i + 1}", exp["Company"], key=f"exp_company_{i}")
        exp["Job Title"] = st.text_input(f"Job Title {i + 1}", exp["Job Title"], key=f"exp_title_{i}")
        exp["Duration"] = st.text_input(f"Duration {i + 1}", exp["Duration"], key=f"exp_duration_{i}")
        exp["Responsibilities"] = st.text_area(f"Responsibilities {i + 1}", exp["Responsibilities"], key=f"exp_responsibilities_{i}")

st.subheader("Projects")
if st.button("+ Add Project"):
    st.session_state.project_entries.append({"Project Name": "", "Description": "", "Duration": ""})

for i, proj in enumerate(st.session_state.project_entries):
    with st.expander(f"Project {i + 1}"):
        proj["Project Name"] = st.text_input(f"Project Name {i + 1}", proj["Project Name"], key=f"proj_name_{i}")
        proj["Description"] = st.text_area(f"Description {i + 1}", proj["Description"], key=f"proj_desc_{i}")
        proj["Duration"] = st.text_input(f"Duration {i + 1}", proj["Duration"], key=f"proj_duration_{i}")

# 🟢 Resume Generation
if submit_button:
    required_inputs = [name, position, email, phone, skills, job_role, job_responsibilities]

    if validate_input(required_inputs):
        # Convert education, experience, and projects into formatted text
        education_text = "\n".join([f"{edu['Institution']} | {edu['Course']} | CGPA: {edu['CGPA']} | {edu['Duration']}" for edu in st.session_state.education_entries]) or "Not Provided"
        experience_text = "\n".join([f"{exp['Company']} | {exp['Job Title']} ({exp['Duration']}): {exp['Responsibilities']}" for exp in st.session_state.experience_entries]) or "No experience listed"
        projects_text = "\n".join([f"{proj['Project Name']} ({proj['Duration']}): {proj['Description']}" for proj in st.session_state.project_entries]) or "No projects listed"

        resume = Resume(
            name=name,
            contact=f"{email} | {phone} | {linkedin or 'N/A'} | {github or 'N/A'} | {location}",
            education=education_text,
            skills=skills,
            experience=experience_text,
            projects=projects_text,
            job=job_role
        )

        st.info("📤 Sending request to AI model...")
        api_response = resume.run()  # Call API

        # 🔹 Display API Response in Streamlit
        st.subheader("🔍 API Response:")
        st.text_area("Raw API Response", api_response, height=200)

        # Check if the response is valid
        if api_response.startswith("⚠"):
            st.error(f"❌ Error: {api_response}")
        else:
            st.text_area("Generated Resume", api_response, height=300)

            # 🔹 Save to DOCX with user name
            # docx_path = f"{name}_resume.docx"
            # doc = Document()
            # doc.add_paragraph(api_response)
            # doc.save(docx_path)

            # st.success(f"✅ Resume generated successfully!")
            # st.success(f"📄 {docx_path} is saved.")
            # st.markdown(generate_download_link(docx_path), unsafe_allow_html=True)

            # 🔹 Save to DOCX with user name
            if "**Expected Output:**" in api_response:
                _, extracted_resume = api_response.split("**Expected Output:**", 1)
                extracted_resume = extracted_resume.strip()
            else:
                extracted_resume = api_response.strip()

            st.write("📄 Extracted Resume Content:")
            st.code(extracted_resume)

            # 📄 Format and Save DOCX
            docx_path = os.path.join(os.getcwd(), f"{name}_resume.docx")
            doc = Document()

            def format_resume(doc, content):
                lines = content.splitlines()
                for line in lines:
                    try:
                        line = line.strip()
                        if not line:
                            doc.add_paragraph("")
                            continue
                        if line.startswith("### "):
                            para = doc.add_paragraph(line.replace("### ", "").strip())
                            run = para.runs[0]
                            run.bold = True
                            run.font.size = Pt(14)
                        elif line.startswith("**") and line.endswith("**"):
                            para = doc.add_paragraph(line.replace("**", "").strip())
                            run = para.runs[0]
                            run.bold = True
                            run.font.size = Pt(12)
                        elif line.startswith("- "):
                            para = doc.add_paragraph(style="List Bullet")
                            run = para.add_run(line[2:].strip())
                            run.font.size = Pt(10.5)
                        else:
                            para = doc.add_paragraph(line)
                            run = para.runs[0]
                            run.font.size = Pt(10.5)
                    except Exception as e:
                        logging.error(f"⚠ Error formatting line: {line} -> {e}")
                        st.warning(f"⚠ Error formatting line: {line}")

            format_resume(doc, extracted_resume)

            try:
                doc.save(docx_path)
                st.success(f"✅ DOCX file saved: {docx_path}")
                st.markdown(generate_download_link(docx_path), unsafe_allow_html=True)
            except Exception as e:
                st.error(f"❌ Failed to save DOCX: {e}")
            
            

    else:
        st.warning("Please fill in all required fields.")
